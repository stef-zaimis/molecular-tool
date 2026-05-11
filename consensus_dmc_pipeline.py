from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

FASTA_PATH = "/home/spinoulis/Desktop/Platygastroidea/InternshipNaturalis2025/REWORKING THE DATA/forUITOTO/Leptacis_allSequences-BOLD-09March2026_aln.fasta"

# Full-alignment DMC sites (1-based)
DMC_SITES = [241, 251, 252, 257, 269, 312, 321, 322,
             376, 406, 449, 511, 514, 561, 578, 580]

# Outputs
OUT_FASTA_TRIMMED = "tipulae_consensus_trimmed.fasta"
OUT_FASTA_UNTRIMMED = "tipulae_consensus_untrimmed.fasta"
OUT_REPORT = "tipulae_consensus_report.txt"
OUT_DOCX = "trimmed_dmc_sequence.docx"

COLUMN_THRESHOLD = 0.5   # keep site if >50% non-gap

IUPAC_TO_SET = {
    "A": {"A"}, "C": {"C"}, "G": {"G"}, "T": {"T"},
    "R": {"A", "G"}, "Y": {"C", "T"},
    "S": {"G", "C"}, "W": {"A", "T"},
    "M": {"A", "C"}, "K": {"G", "T"},
    "B": {"C", "G", "T"}, "D": {"A", "G", "T"},
    "H": {"A", "C", "T"}, "V": {"A", "C", "G"},
    "N": {"A", "C", "G", "T"},
}

SET_TO_IUPAC = {
    frozenset({"A"}): "A",
    frozenset({"C"}): "C",
    frozenset({"G"}): "G",
    frozenset({"T"}): "T",
    frozenset({"A", "G"}): "R",
    frozenset({"C", "T"}): "Y",
    frozenset({"G", "C"}): "S",
    frozenset({"A", "T"}): "W",
    frozenset({"A", "C"}): "M",
    frozenset({"G", "T"}): "K",
    frozenset({"C", "G", "T"}): "B",
    frozenset({"A", "G", "T"}): "D",
    frozenset({"A", "C", "T"}): "H",
    frozenset({"A", "C", "G"}): "V",
    frozenset({"A", "C", "G", "T"}): "N",
}

BASE_COLORS = {
    "A": RGBColor(0x00, 0x80, 0x00),  # green
    "C": RGBColor(0x00, 0x00, 0xFF),  # blue
    "T": RGBColor(0xFF, 0x00, 0x00),  # red
    "G": RGBColor(0xFF, 0x8C, 0x00),  # orange
}


def parse_fasta(path):
    sequences = {}
    with open(path) as f:
        header = None
        seq_chunks = []

        for line in f:
            line = line.strip()
            if line.startswith(">"):
                if header:
                    sequences[header] = "".join(seq_chunks).upper()
                header = line[1:]
                seq_chunks = []
            else:
                seq_chunks.append(line)

        if header:
            sequences[header] = "".join(seq_chunks).upper()

    return sequences


def ungap(seq):
    return seq.replace("-", "")


def write_fasta(header, seq, out_path):
    with open(out_path, "w") as f:
        f.write(f">{header}\n")
        for i in range(0, len(seq), 80):
            f.write(seq[i:i+80] + "\n")


def assign_consensus_state(non_gap_column):
    """
    Ignore gaps.
    Count only A/C/G/T.
    Assign a single base only if f1 - f2 > 0.5.
    Otherwise assign the exact IUPAC code matching the observed base set.
    """
    counts = {}
    for c in non_gap_column:
        if c in "ACGT":
            counts[c] = counts.get(c, 0) + 1

    if not counts:
        return "N"

    total = sum(counts.values())
    ranked = sorted(counts.items(), key=lambda x: (-x[1], x[0]))

    if len(ranked) == 1:
        return ranked[0][0]

    b1, c1 = ranked[0]
    b2, c2 = ranked[1]

    f1 = c1 / total
    f2 = c2 / total

    if (f1 - f2) > 0.5:
        return b1

    base_set = frozenset(counts.keys())
    return SET_TO_IUPAC.get(base_set, "N")


def build_untrimmed_consensus(tip):
    length = len(tip[0])
    consensus = []
    total_N = 0

    for i in range(length):
        column = [s[i] for s in tip]
        non_gap = [c for c in column if c != "-"]

        state = assign_consensus_state(non_gap)
        if state == "N":
            total_N += 1
        consensus.append(state)

    return "".join(consensus), total_N


def build_trimmed_consensus(tip):
    length = len(tip[0])
    n_seq = len(tip)

    consensus = []
    kept_indices = []
    removed_low_cov = 0
    total_N = 0

    for i in range(length):
        column = [s[i] for s in tip]
        non_gap = [c for c in column if c != "-"]
        coverage = len(non_gap) / n_seq

        if coverage <= COLUMN_THRESHOLD:
            removed_low_cov += 1
            continue

        kept_indices.append(i)  # 0-based
        state = assign_consensus_state(non_gap)

        if state == "N":
            total_N += 1

        consensus.append(state)

    return "".join(consensus), kept_indices, removed_low_cov, total_N


def similarity_to_consensus(consensus, seq):
    """
    Compare aligned sequence to untrimmed consensus in alignment coordinates.
    Full match if the sequence base is contained in the consensus code.
    Ignore:
    - consensus N
    - sequence gaps
    - non-ACGT sequence states
    """
    matches = 0
    valid = 0

    for c, s in zip(consensus, seq):
        if c == "N" or s == "-" or s not in "ACGT":
            continue

        valid += 1
        if s in IUPAC_TO_SET.get(c, set()):
            matches += 1

    return matches / valid if valid else 0.0


def cluster_sequences(sequences):
    """
    Group exact duplicates and strict ungapped substrings together.
    Representative = longest ungapped sequence in the group.
    """
    groups = []

    for h, seq in sequences.items():
        s = ungap(seq)
        placed = False

        for g in groups:
            rep = g["rep_seq"]
            if s in rep or rep in s:
                g["members"].append((h, s))
                if len(s) > len(rep):
                    g["rep_seq"] = s
                    g["rep_id"] = h
                placed = True
                break

        if not placed:
            groups.append({
                "rep_id": h,
                "rep_seq": s,
                "members": [(h, s)]
            })

    return groups


def map_dmc_sites_to_trimmed(dmc_sites_1based, kept_indices_0based):
    """
    Map full-alignment DMC positions to trimmed-consensus positions.

    Returns a list of dicts:
    {
        "full_site": 241,
        "kept": True/False,
        "trimmed_pos": 203 or None
    }
    """
    kept_indices_1based = [i + 1 for i in kept_indices_0based]
    mapping = []

    for site in dmc_sites_1based:
        if site in kept_indices_1based:
            trimmed_pos = kept_indices_1based.index(site) + 1
            mapping.append({
                "full_site": site,
                "kept": True,
                "trimmed_pos": trimmed_pos
            })
        else:
            mapping.append({
                "full_site": site,
                "kept": False,
                "trimmed_pos": None
            })

    return mapping


def build_dmc_docx(trimmed_seq, trimmed_positions_1based, output_path):
    """
    Write a Word document containing the trimmed consensus:
    - all lowercase
    - DMC positions uppercase + bold
    - base colours applied
    """
    seq = trimmed_seq.lower()

    doc = Document()

    title = doc.add_paragraph()
    title_run = title.add_run("Trimmed consensus with DMC positions highlighted")
    title_run.bold = True
    title_run.font.size = Pt(12)

    doc.add_paragraph("")

    p = doc.add_paragraph()

    for i, base in enumerate(seq, start=1):
        display_base = base.upper() if i in trimmed_positions_1based else base
        run = p.add_run(display_base)

        run.bold = i in trimmed_positions_1based
        run.font.name = "Courier New"
        run.font.size = Pt(10)

        base_upper = base.upper()
        if base_upper in BASE_COLORS:
            run.font.color.rgb = BASE_COLORS[base_upper]

    doc.save(output_path)


def main():
    sequences = parse_fasta(FASTA_PATH)
    tipulae = {h: s for h, s in sequences.items() if "Leptacis_tipulae" in h}
    tip_list = list(tipulae.values())

    if not tip_list:
        print("No Leptacis_tipulae sequences found.")
        return

    # Build consensuses
    untrimmed, untrim_N = build_untrimmed_consensus(tip_list)
    trimmed, kept_idx, removed_low, trimmed_N = build_trimmed_consensus(tip_list)

    # Write FASTA outputs
    write_fasta("tipulae_consensus_untrimmed", untrimmed, OUT_FASTA_UNTRIMMED)
    write_fasta("tipulae_consensus_trimmed", trimmed, OUT_FASTA_TRIMMED)

    # Group sequences
    groups = cluster_sequences(tipulae)

    # Map DMC sites to trimmed coordinates
    dmc_mapping = map_dmc_sites_to_trimmed(DMC_SITES, kept_idx)
    kept_dmc_positions_trimmed = [x["trimmed_pos"] for x in dmc_mapping if x["kept"]]

    # Write Word document from the actual generated trimmed consensus
    build_dmc_docx(trimmed, set(kept_dmc_positions_trimmed), OUT_DOCX)

    # Write report
    with open(OUT_REPORT, "w") as f:
        f.write("=== CONSENSUS ===\n")
        f.write(f"Input tipulae sequences: {len(tipulae)}\n")
        f.write(f"Original alignment length: {len(tip_list[0])}\n")
        f.write(f"Untrimmed length: {len(untrimmed)}\n")
        f.write(f"Untrimmed N sites: {untrim_N}\n")
        f.write(f"Trimmed length: {len(trimmed)}\n")
        f.write(f"Trimmed N sites: {trimmed_N}\n")
        f.write(f"Removed sites (low coverage): {removed_low}\n\n")

        f.write("=== DMC SITE MAPPING ===\n")
        for item in dmc_mapping:
            if item["kept"]:
                base = trimmed[item["trimmed_pos"] - 1]
                f.write(f"{item['full_site']} -> trimmed {item['trimmed_pos']} -> {base}\n")
            else:
                f.write(f"{item['full_site']} -> DROPPED during trimming\n")
        f.write("\n")

        f.write("=== GROUPS ===\n")
        f.write(f"Number of unique groups: {len(groups)}\n")
        for i, g in enumerate(groups, 1):
            rep_id = g["rep_id"]
            sim = similarity_to_consensus(untrimmed, tipulae[rep_id])

            f.write(f"\nGroup {i}\n")
            f.write(f"Representative: {rep_id}\n")
            f.write(f"Representative length (ungapped): {len(g['rep_seq'])}\n")
            f.write(f"Similarity to untrimmed consensus: {sim:.4f}\n")

            for h, s in sorted(g["members"], key=lambda x: (-len(x[1]), x[0])):
                f.write(f"  - {h} | length={len(s)}\n")

    print("Done.")
    print(f"Wrote: {OUT_FASTA_UNTRIMMED}")
    print(f"Wrote: {OUT_FASTA_TRIMMED}")
    print(f"Wrote: {OUT_REPORT}")
    print(f"Wrote: {OUT_DOCX}")


if __name__ == "__main__":
    main()
